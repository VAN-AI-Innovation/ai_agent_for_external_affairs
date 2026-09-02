from functools import lru_cache
from io import BytesIO
from pathlib import Path
import os
import tempfile

from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.agents.local_contract_analyzer import LocalContractAnalyzer
from app.security.masking import mask_sensitive_text


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".png", ".jpg", ".jpeg"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
PDF_OCR_SCALE = 2.0


class ContractTextExtractionError(ValueError):
    pass


def extract_text_from_contract_document(file_bytes: bytes, filename: str) -> tuple[str, int | None]:
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return _extract_text_from_pdf(file_bytes)
    if extension in {".txt", ".md"}:
        return _extract_text_from_plain_text(file_bytes), None
    if extension == ".docx":
        return _extract_text_from_docx(file_bytes), None
    if extension in IMAGE_EXTENSIONS:
        return _extract_text_from_image(file_bytes, extension), None

    raise ContractTextExtractionError(
        "지원하지 않는 파일 형식입니다. PDF, TXT, MD, DOCX, PNG, JPG, JPEG 파일을 업로드하세요."
    )


def _extract_text_from_pdf(pdf_bytes: bytes) -> tuple[str, int]:
    try:
        reader = PdfReader(BytesIO(pdf_bytes))
    except PdfReadError as exc:
        raise ContractTextExtractionError("PDF 파일을 읽을 수 없습니다.") from exc

    page_texts: dict[int, str] = {}
    scanned_page_numbers: list[int] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        cleaned = text.strip()
        if cleaned:
            page_texts[index] = cleaned
        else:
            scanned_page_numbers.append(index)

    if scanned_page_numbers:
        page_texts.update(_extract_text_from_scanned_pdf_pages(pdf_bytes, scanned_page_numbers))

    full_text = "\n\n".join(
        f"[페이지 {index}]\n{page_texts[index]}"
        for index in range(1, len(reader.pages) + 1)
        if page_texts.get(index)
    ).strip()
    if not full_text:
        raise ContractTextExtractionError(
            "PDF에서 텍스트를 추출하지 못했습니다. 해상도와 글자 선명도를 확인하세요."
        )

    return full_text, len(reader.pages)


def _extract_text_from_scanned_pdf_pages(pdf_bytes: bytes, page_numbers: list[int]) -> dict[int, str]:
    try:
        import fitz
    except ImportError as exc:
        raise ContractTextExtractionError("스캔 PDF OCR을 위해 PyMuPDF 설치가 필요합니다.") from exc

    reader = _get_ocr_reader()
    extracted: dict[int, str] = {}
    temp_paths: list[str] = []

    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as document:
            matrix = fitz.Matrix(PDF_OCR_SCALE, PDF_OCR_SCALE)
            for page_number in page_numbers:
                page = document.load_page(page_number - 1)
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
                    temp_path = temp_file.name
                temp_paths.append(temp_path)
                pixmap.save(temp_path)

                results = reader.readtext(temp_path, detail=0, paragraph=True)
                text = "\n".join(item.strip() for item in results if item.strip()).strip()
                if text:
                    extracted[page_number] = text
    except Exception as exc:
        raise ContractTextExtractionError("스캔 PDF OCR 처리 중 오류가 발생했습니다.") from exc
    finally:
        for temp_path in temp_paths:
            try:
                os.remove(temp_path)
            except OSError:
                pass

    return extracted


def _extract_text_from_plain_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp949"):
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise ContractTextExtractionError("텍스트 파일을 읽을 수 없습니다. UTF-8 또는 CP949 인코딩을 사용하세요.")


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ContractTextExtractionError("DOCX 처리를 위해 python-docx 설치가 필요합니다.") from exc

    document = Document(BytesIO(file_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_cells.append(row_text)

    text = "\n".join([*paragraphs, *table_cells]).strip()
    if not text:
        raise ContractTextExtractionError("DOCX 파일에서 텍스트를 찾지 못했습니다.")
    return text


def _extract_text_from_image(file_bytes: bytes, extension: str) -> str:
    suffix = ".jpg" if extension == ".jpeg" else extension
    temp_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name

        reader = _get_ocr_reader()
        results = reader.readtext(temp_path, detail=0, paragraph=True)
        text = "\n".join(item.strip() for item in results if item.strip()).strip()
    except ImportError as exc:
        raise ContractTextExtractionError("이미지 OCR을 위해 easyocr 설치가 필요합니다.") from exc
    except Exception as exc:
        raise ContractTextExtractionError("이미지에서 텍스트를 추출하지 못했습니다.") from exc
    finally:
        if temp_path:
            try:
                os.remove(temp_path)
            except OSError:
                pass

    if not text:
        raise ContractTextExtractionError("이미지에서 텍스트를 찾지 못했습니다. 해상도와 글자 선명도를 확인하세요.")
    return text


@lru_cache(maxsize=1)
def _get_ocr_reader():
    import easyocr

    return easyocr.Reader(["ko", "en"], gpu=False)


class ContractDocumentAgent:
    def __init__(self):
        self._analyzer = LocalContractAnalyzer()

    def analyze_document(self, file_bytes: bytes, filename: str, review_focus: str | None = None) -> str:
        contract_text, page_count = extract_text_from_contract_document(file_bytes, filename)
        masked_text = mask_sensitive_text(contract_text)
        masked_focus = mask_sensitive_text(review_focus)
        masked_filename = mask_sensitive_text(filename)
        return self._analyzer.analyze(masked_text, masked_filename, page_count or 1, masked_focus)
